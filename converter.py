import pandas as pd
from docx import Document
from docx.shared import Pt
import argparse


def excel_to_word(excel_path, word_path):
    """
    Converts an Excel file to a Word document.

    Each column in the Excel file is placed on a separate page in the Word document.
    On each page, a 2-cell table is created:
        - The first cell contains the column title.
        - The second cell contains the combined content of all cells in that column,
          with each Excel cell's content as a separate paragraph.

    Args:
        excel_path (str): Path to the input Excel file.
        word_path (str): Path to the output Word file.
    """
    # Read the Excel file
    try:
        df = pd.read_excel(excel_path)
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return

    # Create a new Word Document
    doc = Document()

    # Define table style (optional)
    table_style = 'Light List Accent 1'

    # Iterate over each column in the DataFrame
    for idx, column in enumerate(df.columns):
        # Add a page break before each new column except the first
        if idx != 0:
            doc.add_page_break()

        # Add a table with 1 row and 2 columns
        table = doc.add_table(rows=2, cols=1)
        table.style = table_style

        # Access the first row
        row = table.rows[0]

        # First cell: Column Title
        title_cell = row.cells[0]
        title_cell.text = str(column)

        # Format the title (e.g., bold and larger font)
        for paragraph in title_cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(14)

        # Second cell: Combined Content
        row = table.rows[1]
        content_cell = row.cells[0]
        for item in df[column].dropna():
            paragraph = content_cell.add_paragraph(str(item))
            paragraph.style = 'Normal'
            paragraph.paragraph_format.space_after = Pt(6)

        # Optionally, adjust column widths
        # You can set desired widths here if needed

    # Save the Word document
    try:
        doc.save(word_path)
        print(f"Successfully converted '{excel_path}' to '{word_path}'.")
    except Exception as e:
        print(f"Error saving Word file: {e}")


def parse_arguments():
    """
    Parses command-line arguments.

    Returns:
        argparse.Namespace: Parsed arguments containing Excel and Word file paths.
    """
    parser = argparse.ArgumentParser(
        description='Convert Excel columns to a formatted Word document.')
    parser.add_argument('excel_file', help='Path to the input Excel file.')
    parser.add_argument('word_file', help='Path to the output Word file.')
    return parser.parse_args()


def main():
    args = parse_arguments()
    excel_to_word(args.excel_file, args.word_file)


if __name__ == '__main__':
    main()
